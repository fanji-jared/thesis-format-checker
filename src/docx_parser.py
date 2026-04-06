import os
import re
import time
from typing import Optional, List, Tuple, Callable, Dict, Any
from dataclasses import dataclass
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import lxml.etree as etree

from .models import (
    DocumentInfo, ParagraphInfo, ParagraphType, ParagraphFormat,
    FontInfo, RunInfo, PageSettings, SectionInfo, TOCInfo, TOCEntry,
    FigureInfo, TableInfo, FormulaInfo, ReferenceInfo, CitationInfo,
    ChapterInfo, SubSectionInfo, HeadingLevel, ParseResult
)
from .cache import DocumentCache
from .logger import get_logger

logger = get_logger()


class ParseProgress:
    def __init__(self, stage: str, progress: float, message: str, details: Optional[Dict[str, Any]] = None):
        self.stage = stage
        self.progress = progress
        self.message = message
        self.details = details or {}
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            'stage': self.stage,
            'progress': self.progress,
            'message': self.message,
            'details': self.details
        }


class ParseError(Exception):
    def __init__(self, message: str, error_type: str = "unknown", details: Optional[Dict] = None):
        super().__init__(message)
        self.error_type = error_type
        self.details = details or {}
        self.user_message = self._get_user_message()
    
    def _get_user_message(self) -> str:
        error_messages = {
            'file_not_found': '文档文件不存在，请检查文件路径是否正确。',
            'file_corrupted': '文档文件可能已损坏，请尝试用Word打开并修复。',
            'invalid_format': '文档格式无效，请确保文件是有效的.docx格式。',
            'permission_denied': '无法访问文档文件，请检查文件是否被其他程序占用。',
            'memory_error': '文档太大，内存不足。请尝试关闭其他程序后重试。',
            'encoding_error': '文档编码异常，请尝试另存为新文件后重试。',
            'unknown': f'文档解析失败: {str(self)}'
        }
        return error_messages.get(self.error_type, error_messages['unknown'])


class DocxParser:
    HEADING_1_PATTERN = re.compile(r'^第\s*(\d+)\s*章\s+(.+)$')
    HEADING_2_PATTERN = re.compile(r'^(\d+)\.(\d+)\s+(.+)$')
    HEADING_3_PATTERN = re.compile(r'^(\d+)\.(\d+)\.(\d+)\s+(.+)$')
    FIGURE_CAPTION_PATTERN = re.compile(r'^图\s*(\d+)-(\d+)\s*(.*)$')
    TABLE_CAPTION_PATTERN = re.compile(r'^表\s*(\d+)-(\d+)\s*(.*)$')
    FORMULA_PATTERN = re.compile(r'\((\d+)-(\d+)\)$')
    CITATION_PATTERN = re.compile(r'\[(\d+)\]')
    REFERENCE_START_PATTERN = re.compile(r'^参考文献\s*$|^References\s*$', re.IGNORECASE)
    
    _cache: Optional[DocumentCache] = None

    def __init__(self, file_path: Optional[str] = None, use_cache: bool = True, progress_callback: Optional[Callable[[ParseProgress], None]] = None):
        self.file_path = file_path
        self.document: Optional[Document] = None
        self._paragraph_index = 0
        self._current_chapter = 0
        self._figure_counter = {}
        self._table_counter = {}
        self._formula_counter = {}
        self._use_cache = use_cache
        self._progress_callback = progress_callback
        
        if use_cache and DocxParser._cache is None:
            DocxParser._cache = DocumentCache()

    def _update_progress(self, stage: str, progress: float, message: str, details: Optional[Dict[str, Any]] = None):
        if self._progress_callback:
            progress_info = ParseProgress(stage, progress, message, details)
            self._progress_callback(progress_info)
        logger.debug(f"[{stage}] {progress*100:.0f}% - {message}")

    def parse(self, file_path: Optional[str] = None) -> ParseResult:
        if file_path:
            self.file_path = file_path
        
        if not self.file_path:
            error = ParseError("未指定文件路径", "file_not_found")
            logger.error(error.user_message)
            return ParseResult(success=False, error_message=error.user_message)
        
        start_time = time.time()
        
        if not os.path.exists(self.file_path):
            error = ParseError(f"文件不存在: {self.file_path}", "file_not_found")
            logger.error(error.user_message)
            return ParseResult(success=False, error_message=error.user_message)
        
        if self._use_cache and DocxParser._cache:
            cached_result = DocxParser._cache.get(self.file_path)
            if cached_result:
                logger.info(f"从缓存加载文档: {self.file_path}")
                self._update_progress("completed", 1.0, "从缓存加载完成")
                return ParseResult(success=True, document_info=cached_result)
        
        try:
            self._update_progress("loading", 0.0, "正在加载文档...")
            logger.info(f"开始解析文档: {self.file_path}")
            
            self.document = Document(self.file_path)
            
            self._update_progress("parsing", 0.1, "正在解析文档结构...")
            doc_info = self._extract_document_info()
            
            if self._use_cache and DocxParser._cache:
                DocxParser._cache.set(self.file_path, doc_info)
                logger.info("文档已缓存")
            
            elapsed = time.time() - start_time
            self._update_progress("completed", 1.0, f"解析完成，耗时 {elapsed:.2f}秒")
            logger.info(f"文档解析完成，耗时 {elapsed:.2f}秒")
            
            return ParseResult(success=True, document_info=doc_info)
            
        except FileNotFoundError as e:
            error = ParseError(str(e), "file_not_found")
            logger.error(error.user_message)
            return ParseResult(success=False, error_message=error.user_message)
        except PermissionError as e:
            error = ParseError(str(e), "permission_denied")
            logger.error(error.user_message)
            return ParseResult(success=False, error_message=error.user_message)
        except MemoryError as e:
            error = ParseError(str(e), "memory_error")
            logger.error(error.user_message)
            return ParseResult(success=False, error_message=error.user_message)
        except Exception as e:
            error_msg = str(e)
            if "corrupted" in error_msg.lower() or "invalid" in error_msg.lower():
                error = ParseError(error_msg, "file_corrupted")
            else:
                error = ParseError(error_msg, "unknown")
            logger.error(f"文档解析失败: {error_msg}")
            return ParseResult(success=False, error_message=error.user_message)

    def _extract_document_info(self) -> DocumentInfo:
        file_name = os.path.basename(self.file_path)
        doc_info = DocumentInfo(
            file_path=self.file_path,
            file_name=file_name
        )

        self._update_progress("parsing", 0.15, "正在提取页面设置...")
        doc_info.page_settings = self._extract_page_settings()
        doc_info.sections = self._extract_sections()
        
        self._update_progress("parsing", 0.25, "正在提取段落...")
        doc_info.all_paragraphs = self._extract_all_paragraphs()
        doc_info.paragraph_count = len(doc_info.all_paragraphs)
        doc_info.section_count = len(doc_info.sections)
        doc_info.character_count = sum(len(p.text) for p in doc_info.all_paragraphs)

        self._update_progress("parsing", 0.35, "正在提取目录...")
        doc_info.toc_info = self._extract_toc(doc_info.all_paragraphs)
        doc_info.has_toc = doc_info.toc_info is not None

        self._update_progress("parsing", 0.45, "正在识别章节...")
        self._identify_headings(doc_info.all_paragraphs)
        doc_info.chapters = self._extract_chapters(doc_info.all_paragraphs)

        self._update_progress("parsing", 0.55, "正在提取图片...")
        doc_info.figures = self._extract_figures(doc_info.all_paragraphs)
        
        self._update_progress("parsing", 0.65, "正在提取表格...")
        doc_info.tables = self._extract_tables(doc_info.all_paragraphs)
        
        self._update_progress("parsing", 0.75, "正在提取公式...")
        doc_info.formulas = self._extract_formulas(doc_info.all_paragraphs)

        self._update_progress("parsing", 0.85, "正在提取参考文献...")
        ref_result = self._extract_references(doc_info.all_paragraphs)
        doc_info.references = ref_result['references']
        doc_info.citations = ref_result['citations']
        doc_info.reference_section_start = ref_result['start_index']
        doc_info.reference_section_end = ref_result['end_index']

        self._update_progress("parsing", 0.95, "正在计算页数...")
        doc_info.page_count = self._estimate_page_count(doc_info)

        return doc_info

    def _extract_page_settings(self) -> List[PageSettings]:
        settings_list = []
        for section in self.document.sections:
            settings = PageSettings(
                top_margin=self._emu_to_cm(section.top_margin) if section.top_margin else 0,
                bottom_margin=self._emu_to_cm(section.bottom_margin) if section.bottom_margin else 0,
                left_margin=self._emu_to_cm(section.left_margin) if section.left_margin else 0,
                right_margin=self._emu_to_cm(section.right_margin) if section.right_margin else 0,
                gutter=self._emu_to_cm(section.gutter) if hasattr(section, 'gutter') and section.gutter else 0,
                gutter_position=self._get_gutter_position(section),
                header_distance=self._emu_to_cm(section.header_distance) if section.header_distance else 0,
                footer_distance=self._emu_to_cm(section.footer_distance) if section.footer_distance else 0,
                page_width=self._emu_to_cm(section.page_width) if section.page_width else 0,
                page_height=self._emu_to_cm(section.page_height) if section.page_height else 0,
                orientation=self._get_orientation(section)
            )
            settings_list.append(settings)
        return settings_list

    def _extract_sections(self) -> List[SectionInfo]:
        sections = []
        for idx, section in enumerate(self.document.sections):
            settings = PageSettings(
                top_margin=self._emu_to_cm(section.top_margin) if section.top_margin else 0,
                bottom_margin=self._emu_to_cm(section.bottom_margin) if section.bottom_margin else 0,
                left_margin=self._emu_to_cm(section.left_margin) if section.left_margin else 0,
                right_margin=self._emu_to_cm(section.right_margin) if section.right_margin else 0,
                gutter=self._emu_to_cm(section.gutter) if hasattr(section, 'gutter') and section.gutter else 0,
                gutter_position=self._get_gutter_position(section),
                header_distance=self._emu_to_cm(section.header_distance) if section.header_distance else 0,
                footer_distance=self._emu_to_cm(section.footer_distance) if section.footer_distance else 0,
                page_width=self._emu_to_cm(section.page_width) if section.page_width else 0,
                page_height=self._emu_to_cm(section.page_height) if section.page_height else 0,
                orientation=self._get_orientation(section)
            )
            section_info = SectionInfo(
                section_index=idx,
                start_page=1,
                page_settings=settings,
                has_different_first_page=self._has_different_first_page(section),
                has_different_odd_even=self._has_different_odd_even(section)
            )
            sections.append(section_info)
        return sections

    def _extract_all_paragraphs(self) -> List[ParagraphInfo]:
        paragraphs = []
        self._paragraph_index = 0

        for para in self.document.paragraphs:
            para_info = self._extract_paragraph_info(para, self._paragraph_index)
            paragraphs.append(para_info)
            self._paragraph_index += 1

        return paragraphs

    def _extract_paragraph_info(self, para, index: int) -> ParagraphInfo:
        text = para.text.strip()
        runs = self._extract_runs(para)
        para_format = self._extract_paragraph_format(para)
        style_name = para.style.name if para.style else None

        para_type = self._identify_paragraph_type(text, style_name, runs)
        heading_level = self._get_heading_level(style_name, text)

        para_info = ParagraphInfo(
            index=index,
            text=text,
            paragraph_type=para_type,
            runs=runs,
            format=para_format,
            style_name=style_name,
            heading_level=heading_level
        )

        if heading_level != HeadingLevel.NONE:
            self._extract_heading_info(para_info, text)

        return para_info

    def _extract_runs(self, para) -> List[RunInfo]:
        runs = []
        for run in para.runs:
            font_info = self._extract_font_info(run)
            run_info = RunInfo(
                text=run.text,
                font=font_info,
                is_formula=self._is_formula_run(run),
                is_superscript=run.font.superscript if run.font else False,
                is_subscript=run.font.subscript if run.font else False
            )
            runs.append(run_info)
        return runs

    def _extract_font_info(self, run) -> FontInfo:
        font = run.font
        font_info = FontInfo(name="")

        if font:
            font_info.name = font.name or ""
            font_info.size = self._pt_to_pt(font.size) if font.size else None
            font_info.bold = font.bold or False
            font_info.italic = font.italic or False
            font_info.underline = font.underline or False

            if hasattr(font, 'color') and font.color and font.color.rgb:
                font_info.color = str(font.color.rgb)

            rPr = run._element.rPr
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    font_info.name_ascii = rFonts.get(qn('w:ascii'))
                    font_info.name_east_asia = rFonts.get(qn('w:eastAsia'))

        return font_info

    def _extract_paragraph_format(self, para) -> ParagraphFormat:
        pf = para.paragraph_format
        para_format = ParagraphFormat()

        if pf:
            para_format.alignment = self._get_alignment(pf.alignment)
            para_format.first_line_indent = self._emu_to_cm(pf.first_line_indent) if pf.first_line_indent else None
            para_format.left_indent = self._emu_to_cm(pf.left_indent) if pf.left_indent else None
            para_format.right_indent = self._emu_to_cm(pf.right_indent) if pf.right_indent else None
            if hasattr(pf, 'hanging_indent'):
                para_format.hanging_indent = self._emu_to_cm(pf.hanging_indent) if pf.hanging_indent else None
            para_format.line_spacing = pf.line_spacing
            para_format.line_spacing_rule = self._get_line_spacing_rule(pf.line_spacing_rule)
            para_format.space_before = self._pt_to_pt(pf.space_before) if pf.space_before else None
            para_format.space_after = self._pt_to_pt(pf.space_after) if pf.space_after else None

        pPr = para._element.pPr
        if pPr is not None:
            outlineLvl = pPr.find(qn('w:outlineLvl'))
            if outlineLvl is not None:
                para_format.outline_level = int(outlineLvl.get(qn('w:val'), 0))

        return para_format

    def _identify_paragraph_type(self, text: str, style_name: Optional[str], runs: List[RunInfo]) -> ParagraphType:
        if not text:
            return ParagraphType.UNKNOWN

        if self.HEADING_1_PATTERN.match(text):
            return ParagraphType.HEADING_1
        if self.HEADING_2_PATTERN.match(text):
            return ParagraphType.HEADING_2
        if self.HEADING_3_PATTERN.match(text):
            return ParagraphType.HEADING_3

        if self.FIGURE_CAPTION_PATTERN.match(text):
            return ParagraphType.FIGURE_CAPTION
        if self.TABLE_CAPTION_PATTERN.match(text):
            return ParagraphType.TABLE_CAPTION

        if style_name:
            style_lower = style_name.lower()
            if 'heading 1' in style_lower or '标题 1' in style_name:
                return ParagraphType.HEADING_1
            if 'heading 2' in style_lower or '标题 2' in style_name:
                return ParagraphType.HEADING_2
            if 'heading 3' in style_lower or '标题 3' in style_name:
                return ParagraphType.HEADING_3
            if 'toc' in style_lower or '目录' in style_name:
                return ParagraphType.TOC_ENTRY
            if 'caption' in style_lower or '题注' in style_name:
                if '图' in text:
                    return ParagraphType.FIGURE_CAPTION
                if '表' in text:
                    return ParagraphType.TABLE_CAPTION

        if self.REFERENCE_START_PATTERN.match(text):
            return ParagraphType.HEADING_1

        if self._has_formula_content(runs):
            return ParagraphType.FORMULA

        return ParagraphType.BODY_TEXT

    def _get_heading_level(self, style_name: Optional[str], text: str) -> HeadingLevel:
        if self.HEADING_1_PATTERN.match(text) or self.REFERENCE_START_PATTERN.match(text):
            return HeadingLevel.LEVEL_1
        if self.HEADING_2_PATTERN.match(text):
            return HeadingLevel.LEVEL_2
        if self.HEADING_3_PATTERN.match(text):
            return HeadingLevel.LEVEL_3

        if style_name:
            style_lower = style_name.lower()
            if 'heading 1' in style_lower or '标题 1' in style_name:
                return HeadingLevel.LEVEL_1
            if 'heading 2' in style_lower or '标题 2' in style_name:
                return HeadingLevel.LEVEL_2
            if 'heading 3' in style_lower or '标题 3' in style_name:
                return HeadingLevel.LEVEL_3

        return HeadingLevel.NONE

    def _extract_heading_info(self, para_info: ParagraphInfo, text: str):
        match1 = self.HEADING_1_PATTERN.match(text)
        if match1:
            para_info.chapter_number = int(match1.group(1))
            para_info.heading_number = f"第{match1.group(1)}章"
            para_info.heading_title = match1.group(2).strip()
            return

        if self.REFERENCE_START_PATTERN.match(text):
            para_info.heading_number = "参考文献"
            para_info.heading_title = "参考文献"
            return

        match2 = self.HEADING_2_PATTERN.match(text)
        if match2:
            para_info.chapter_number = int(match2.group(1))
            para_info.section_number = f"{match2.group(1)}.{match2.group(2)}"
            para_info.heading_number = para_info.section_number
            para_info.heading_title = match2.group(3).strip()
            return

        match3 = self.HEADING_3_PATTERN.match(text)
        if match3:
            para_info.chapter_number = int(match3.group(1))
            para_info.section_number = f"{match3.group(1)}.{match3.group(2)}.{match3.group(3)}"
            para_info.heading_number = para_info.section_number
            para_info.heading_title = match3.group(4).strip()

    def _extract_toc(self, paragraphs: List[ParagraphInfo]) -> Optional[TOCInfo]:
        toc_entries = []
        in_toc = False
        toc_title = "目录"

        for para in paragraphs:
            text = para.text.strip()
            
            if text == "目录" or text.lower() == "contents":
                in_toc = True
                toc_title = text
                continue

            if in_toc:
                if para.heading_level != HeadingLevel.NONE and not para.style_name or 'toc' not in (para.style_name or '').lower():
                    if self.HEADING_1_PATTERN.match(text) or self.HEADING_2_PATTERN.match(text):
                        break

                if text and para.style_name and ('toc' in para.style_name.lower() or '目录' in para.style_name):
                    level = self._get_toc_level(para.style_name)
                    entry = self._parse_toc_entry(text, level)
                    toc_entries.append(entry)

        if toc_entries:
            is_auto = self._check_auto_toc()
            return TOCInfo(entries=toc_entries, is_auto_generated=is_auto, title=toc_title)

        return None

    def _get_toc_level(self, style_name: str) -> int:
        if 'toc 1' in style_name.lower() or '目录 1' in style_name:
            return 1
        if 'toc 2' in style_name.lower() or '目录 2' in style_name:
            return 2
        if 'toc 3' in style_name.lower() or '目录 3' in style_name:
            return 3
        return 1

    def _parse_toc_entry(self, text: str, level: int) -> TOCEntry:
        heading_number = None
        heading_title = text

        match1 = self.HEADING_1_PATTERN.match(text)
        if match1:
            heading_number = f"第{match1.group(1)}章"
            heading_title = match1.group(2).strip()
        else:
            match2 = self.HEADING_2_PATTERN.match(text)
            if match2:
                heading_number = f"{match2.group(1)}.{match2.group(2)}"
                heading_title = match2.group(3).strip()
            else:
                match3 = self.HEADING_3_PATTERN.match(text)
                if match3:
                    heading_number = f"{match3.group(1)}.{match3.group(2)}.{match3.group(3)}"
                    heading_title = match3.group(4).strip()

        return TOCEntry(
            level=level,
            text=text,
            heading_number=heading_number,
            heading_title=heading_title
        )

    def _check_auto_toc(self) -> bool:
        for para in self.document.paragraphs:
            fldChar = para._element.findall('.//' + qn('w:fldChar'))
            if fldChar:
                for fc in fldChar:
                    if fc.get(qn('w:fldCharType')) == 'begin':
                        instrText = para._element.find('.//' + qn('w:instrText'))
                        if instrText is not None and 'TOC' in (instrText.text or ''):
                            return True
        return False

    def _identify_headings(self, paragraphs: List[ParagraphInfo]):
        current_chapter = 0
        for para in paragraphs:
            if para.heading_level == HeadingLevel.LEVEL_1:
                if para.chapter_number:
                    current_chapter = para.chapter_number
            elif para.heading_level in (HeadingLevel.LEVEL_2, HeadingLevel.LEVEL_3):
                if current_chapter > 0 and para.chapter_number is None:
                    para.chapter_number = current_chapter

    def _extract_chapters(self, paragraphs: List[ParagraphInfo]) -> List[ChapterInfo]:
        chapters = []
        current_chapter = None
        current_section = None
        chapter_start_idx = 0

        for idx, para in enumerate(paragraphs):
            if para.heading_level == HeadingLevel.LEVEL_1:
                if current_chapter:
                    current_chapter.end_paragraph_index = idx - 1
                    if current_section:
                        current_section.end_paragraph_index = idx - 1
                    chapters.append(current_chapter)

                chapter_start_idx = idx
                current_chapter = ChapterInfo(
                    chapter_number=para.chapter_number or 0,
                    chapter_title=para.heading_title or "",
                    heading_paragraph=para,
                    start_paragraph_index=idx
                )
                current_section = None

            elif para.heading_level == HeadingLevel.LEVEL_2 and current_chapter:
                if current_section:
                    current_section.end_paragraph_index = idx - 1
                    current_chapter.sections.append(current_section)

                current_section = SubSectionInfo(
                    section_number=para.section_number or "",
                    section_title=para.heading_title or "",
                    heading_paragraph=para,
                    level=2,
                    start_paragraph_index=idx
                )

            elif para.heading_level == HeadingLevel.LEVEL_3 and current_section:
                sub_section = SubSectionInfo(
                    section_number=para.section_number or "",
                    section_title=para.heading_title or "",
                    heading_paragraph=para,
                    level=3,
                    start_paragraph_index=idx,
                    end_paragraph_index=idx
                )
                current_section.sub_sections.append(sub_section)

        if current_chapter:
            current_chapter.end_paragraph_index = len(paragraphs) - 1
            if current_section:
                current_section.end_paragraph_index = len(paragraphs) - 1
                current_chapter.sections.append(current_section)
            chapters.append(current_chapter)

        return chapters

    def _extract_figures(self, paragraphs: List[ParagraphInfo]) -> List[FigureInfo]:
        figures = []
        figure_idx = 0
        current_chapter = 0

        for para in paragraphs:
            if para.heading_level == HeadingLevel.LEVEL_1 and para.chapter_number:
                current_chapter = para.chapter_number

            if para.paragraph_type == ParagraphType.FIGURE_CAPTION:
                match = self.FIGURE_CAPTION_PATTERN.match(para.text)
                if match:
                    chapter_num = int(match.group(1))
                    fig_num = int(match.group(2))
                    caption_text = match.group(3).strip()

                    figure = FigureInfo(
                        index=figure_idx,
                        figure_id=f"图{chapter_num}-{fig_num}",
                        caption=caption_text,
                        caption_paragraph=para,
                        position_paragraph_index=para.index,
                        chapter_number=chapter_num,
                        figure_number=fig_num
                    )
                    figures.append(figure)
                    figure_idx += 1

        return figures

    def _extract_tables(self, paragraphs: List[ParagraphInfo]) -> List[TableInfo]:
        tables = []
        table_idx = 0

        for table in self.document.tables:
            table_info = self._extract_table_info(table, table_idx, paragraphs)
            tables.append(table_info)
            table_idx += 1

        return tables

    def _extract_table_info(self, table, index: int, paragraphs: List[ParagraphInfo]) -> TableInfo:
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0

        cell_contents = []
        for row in table.rows:
            row_content = []
            for cell in row.cells:
                row_content.append(cell.text.strip())
            cell_contents.append(row_content)

        borders = self._analyze_table_borders(table)
        is_three_line = self._is_three_line_table(borders, rows)

        caption = None
        caption_para = None
        table_id = None
        chapter_num = None
        table_num = None

        for para in paragraphs:
            if para.paragraph_type == ParagraphType.TABLE_CAPTION:
                match = self.TABLE_CAPTION_PATTERN.match(para.text)
                if match:
                    caption = match.group(3).strip()
                    caption_para = para
                    chapter_num = int(match.group(1))
                    table_num = int(match.group(2))
                    table_id = f"表{chapter_num}-{table_num}"
                    break

        return TableInfo(
            index=index,
            table_id=table_id,
            caption=caption,
            caption_paragraph=caption_para,
            rows=rows,
            columns=cols,
            is_three_line_table=is_three_line,
            has_borders=borders,
            cell_contents=cell_contents,
            chapter_number=chapter_num,
            table_number=table_num
        )

    def _analyze_table_borders(self, table) -> dict:
        borders = {
            'top': False,
            'bottom': False,
            'left': False,
            'right': False,
            'inside_h': False,
            'inside_v': False
        }

        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is not None:
                tblBorders = tblPr.find(qn('w:tblBorders'))
                if tblBorders is not None:
                    for border_name in borders.keys():
                        border = tblBorders.find(qn(f'w:{border_name}'))
                        if border is not None:
                            val = border.get(qn('w:val'))
                            borders[border_name] = val is not None and val != 'nil'
        except Exception:
            pass

        return borders

    def _is_three_line_table(self, borders: dict, rows: int) -> bool:
        if borders.get('left', True) or borders.get('right', True):
            return False

        if borders.get('top') and borders.get('bottom'):
            if not borders.get('inside_v'):
                return True

        return False

    def _extract_formulas(self, paragraphs: List[ParagraphInfo]) -> List[FormulaInfo]:
        formulas = []
        formula_idx = 0
        current_chapter = 0

        for para in paragraphs:
            if para.heading_level == HeadingLevel.LEVEL_1 and para.chapter_number:
                current_chapter = para.chapter_number

            if para.paragraph_type == ParagraphType.FORMULA:
                match = self.FORMULA_PATTERN.search(para.text)
                if match:
                    chapter_num = int(match.group(1))
                    formula_num = int(match.group(2))

                    formula = FormulaInfo(
                        index=formula_idx,
                        formula_id=f"({chapter_num}-{formula_num})",
                        formula_text=para.text,
                        formula_paragraph=para,
                        is_omath=self._has_omath(para),
                        position_paragraph_index=para.index,
                        chapter_number=chapter_num,
                        formula_number=formula_num
                    )
                    formulas.append(formula)
                    formula_idx += 1

        return formulas

    def _has_formula_content(self, runs: List[RunInfo]) -> bool:
        for run in runs:
            if run.is_formula:
                return True
        return False

    def _is_formula_run(self, run) -> bool:
        r = run._element
        omath = r.find('.//' + qn('m:oMath'))
        return omath is not None

    def _has_omath(self, para: ParagraphInfo) -> bool:
        for run in para.runs:
            if run.is_formula:
                return True
        return False

    def _extract_references(self, paragraphs: List[ParagraphInfo]) -> dict:
        references = []
        citations = []
        ref_start_idx = None
        ref_end_idx = None
        ref_idx = 0
        in_reference_section = False

        for idx, para in enumerate(paragraphs):
            if para.paragraph_type == ParagraphType.HEADING_1 and self.REFERENCE_START_PATTERN.match(para.text):
                in_reference_section = True
                ref_start_idx = idx
                continue

            if in_reference_section and para.text.strip():
                ref_info = self._parse_reference(para.text, ref_idx, para)
                references.append(ref_info)
                ref_idx += 1
                ref_end_idx = idx

        citations = self._extract_citations(paragraphs, ref_start_idx)

        return {
            'references': references,
            'citations': citations,
            'start_index': ref_start_idx,
            'end_index': ref_end_idx
        }

    def _parse_reference(self, text: str, index: int, para: ParagraphInfo) -> ReferenceInfo:
        ref_info = ReferenceInfo(
            index=index,
            full_text=text,
            paragraph=para
        )

        patterns = {
            'J': r'\[J\]',
            'M': r'\[M\]',
            'D': r'\[D\]',
            'P': r'\[P\]',
            'S': r'\[S\]',
            'R': r'\[R\]',
            'N': r'\[N\]',
            'EB/OL': r'\[EB/OL\]'
        }

        for ref_type, pattern in patterns.items():
            if re.search(pattern, text):
                ref_info.reference_type = ref_type
                break

        author_match = re.match(r'^\[?\d+\]?\s*([^\.]+?)\.', text)
        if author_match:
            ref_info.authors = author_match.group(1).strip()

        year_match = re.search(r',?\s*(\d{4})[a-z]?', text)
        if year_match:
            ref_info.year = year_match.group(1)

        return ref_info

    def _extract_citations(self, paragraphs: List[ParagraphInfo], ref_section_start: Optional[int]) -> List[CitationInfo]:
        citations = []
        seen_citations = set()

        end_idx = ref_section_start if ref_section_start else len(paragraphs)

        for para in paragraphs[:end_idx]:
            for match in self.CITATION_PATTERN.finditer(para.text):
                citation_num = int(match.group(1))
                if citation_num not in seen_citations:
                    is_superscript = self._is_citation_superscript(para, match.start())
                    citation = CitationInfo(
                        citation_number=citation_num,
                        citation_text=match.group(0),
                        paragraph_index=para.index,
                        position_in_paragraph=match.start(),
                        is_superscript=is_superscript
                    )
                    citations.append(citation)
                    seen_citations.add(citation_num)

        return sorted(citations, key=lambda x: x.citation_number)

    def _is_citation_superscript(self, para: ParagraphInfo, position: int) -> bool:
        current_pos = 0
        for run in para.runs:
            run_start = current_pos
            run_end = current_pos + len(run.text)
            if run_start <= position < run_end:
                return run.is_superscript
            current_pos = run_end
        return False

    def _estimate_page_count(self, doc_info: DocumentInfo) -> int:
        if not doc_info.page_settings:
            return 1

        page_height = doc_info.page_settings[0].page_height
        top_margin = doc_info.page_settings[0].top_margin
        bottom_margin = doc_info.page_settings[0].bottom_margin
        header_distance = doc_info.page_settings[0].header_distance
        footer_distance = doc_info.page_settings[0].footer_distance

        if page_height <= 0:
            return 1

        usable_height = page_height - top_margin - bottom_margin - header_distance - footer_distance
        if usable_height <= 0:
            usable_height = page_height * 0.7

        total_text_length = sum(len(p.text) for p in doc_info.all_paragraphs)
        avg_chars_per_line = 40
        line_height_cm = 0.5
        lines_per_page = int(usable_height / line_height_cm) if usable_height > 0 else 40

        estimated_pages = max(1, (total_text_length // (avg_chars_per_line * lines_per_page)) + 1)

        estimated_pages += len(doc_info.figures) // 2
        estimated_pages += len(doc_info.tables) // 2

        return estimated_pages

    def _emu_to_cm(self, value) -> float:
        if value is None:
            return 0.0
        try:
            if hasattr(value, 'emu'):
                return value.cm
            return float(value) / 914400 * 2.54
        except (AttributeError, TypeError):
            return 0.0

    def _pt_to_pt(self, value) -> Optional[float]:
        if value is None:
            return None
        try:
            if hasattr(value, 'pt'):
                return value.pt
            return float(value)
        except (AttributeError, TypeError):
            return None

    def _get_alignment(self, alignment) -> Optional[str]:
        if alignment is None:
            return None
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: 'left',
            WD_ALIGN_PARAGRAPH.CENTER: 'center',
            WD_ALIGN_PARAGRAPH.RIGHT: 'right',
            WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
            WD_ALIGN_PARAGRAPH.DISTRIBUTE: 'distribute'
        }
        return alignment_map.get(alignment)

    def _get_line_spacing_rule(self, rule) -> Optional[str]:
        if rule is None:
            return None
        rule_map = {
            WD_LINE_SPACING.SINGLE: 'single',
            WD_LINE_SPACING.ONE_POINT_FIVE: 'one_point_five',
            WD_LINE_SPACING.DOUBLE: 'double',
            WD_LINE_SPACING.AT_LEAST: 'at_least',
            WD_LINE_SPACING.EXACTLY: 'exactly',
            WD_LINE_SPACING.MULTIPLE: 'multiple'
        }
        return rule_map.get(rule)

    def _get_gutter_position(self, section) -> str:
        try:
            if hasattr(section, 'gutter_position'):
                return str(section.gutter_position)
        except AttributeError:
            pass
        return "left"

    def _get_orientation(self, section) -> str:
        try:
            if section.orientation == WD_ORIENT.LANDSCAPE:
                return 'landscape'
        except AttributeError:
            pass
        return 'portrait'

    def _has_different_first_page(self, section) -> bool:
        try:
            return section.different_first_page_header_footer
        except AttributeError:
            return False

    def _has_different_odd_even(self, section) -> bool:
        try:
            return section.odd_and_even_pages_header_footer
        except AttributeError:
            return False

    def get_document_summary(self) -> dict:
        result = self.parse()
        if not result.success or not result.document_info:
            return {'error': result.error_message}

        doc = result.document_info
        return {
            'file_name': doc.file_name,
            'page_count': doc.page_count,
            'paragraph_count': doc.paragraph_count,
            'section_count': doc.section_count,
            'character_count': doc.character_count,
            'has_toc': doc.has_toc,
            'is_toc_auto_generated': doc.toc_info.is_auto_generated if doc.toc_info else False,
            'chapter_count': len(doc.chapters),
            'figure_count': len(doc.figures),
            'table_count': len(doc.tables),
            'formula_count': len(doc.formulas),
            'reference_count': len(doc.references),
            'citation_count': len(doc.citations)
        }
